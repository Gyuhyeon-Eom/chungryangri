"""회사소개서 첨부용 보고서를 docx로 생성한다.

편집 가능해야 하므로 표는 네이티브 docx 표로, 차트만 이미지로 넣는다.
차트는 SVG로 만든 뒤 rsvg-convert 로 PNG 변환한다.

문체 원칙
  - 서사적 연결어(다만, 결론부터 말하면, 사정이 달라진다)를 쓰지 않는다
  - 표로 표현 가능한 내용은 문장으로 풀지 않는다
  - 각 절은 관측 → 표/차트 → 해석 2문장 이내로 닫는다
  - 대응 방향은 별도 라벨 없이 표의 한 열 또는 절 마지막 문장으로 둔다

    uv run python src/build_docx.py
"""

import subprocess
import tempfile
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

import charts as C
from config import DATA_PROCESSED, ROOT

OUT = ROOT / "output" / "청량리_전통시장_상권분석_보고서.docx"
TMP = Path(tempfile.mkdtemp(prefix="chart_"))

FONT = "맑은 고딕"
INK = RGBColor(0x11, 0x11, 0x11)
ACCENT = RGBColor(0x1F, 0x3F, 0x5C)
MUTED = RGBColor(0x6B, 0x6B, 0x6B)


# ---------------------------------------------------------------- 문서 유틸
def set_font(run, size=10, bold=False, color=INK):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def para(doc, text="", size=9.5, color=INK, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.4
    for i, chunk in enumerate(text.split("**")):
        if chunk:
            set_font(p.add_run(chunk), size, i % 2 == 1, color)
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 11)
    p.paragraph_format.space_after = Pt(5)
    set_font(p.add_run(text), 13 if level == 1 else 11, True, ACCENT)
    if level == 1:
        bd = OxmlElement("w:pBdr")
        bo = OxmlElement("w:bottom")
        bo.set(qn("w:val"), "single")
        bo.set(qn("w:sz"), "12")
        bo.set(qn("w:color"), "1F3F5C")
        bd.append(bo)
        p._element.get_or_add_pPr().append(bd)
    return p


def caption(doc, text):
    """표·그림 캡션. 대상 아래에 가운데 정렬로 둔다."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(10)
    set_font(p.add_run(text), 8.5, False, MUTED)


def note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    set_font(p.add_run(text), 8.5, False, MUTED)


def table(doc, headers, rows, widths=None, align_right=(), align_center=(), cap=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(str(h)), 8.8, True)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "E9EDF1")
        cell._tc.get_or_add_tcPr().append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.25
            if i in align_right:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif i in align_center:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for k, chunk in enumerate(str(v).split("**")):
                if chunk:
                    set_font(p.add_run(chunk), 8.8, k % 2 == 1)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    if cap:
        caption(doc, cap)
    return t


def add_chart(doc, svg: str, name: str, cap=None, width_cm=16.4):
    svg_path, png_path = TMP / f"{name}.svg", TMP / f"{name}.png"
    svg_path.write_text(svg, encoding="utf-8")
    subprocess.run(["rsvg-convert", "-w", "1900", "-b", "white",
                    "-o", str(png_path), str(svg_path)], check=True)
    doc.add_picture(str(png_path), width=Cm(width_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if cap:
        caption(doc, cap)


# ---------------------------------------------------------------- 본문
def build():
    bm = pd.read_csv(DATA_PROCESSED / "seoul_benchmark.csv")
    ss = pd.read_csv(DATA_PROCESSED / "shift_share.csv", index_col=0)
    dt = pd.read_csv(DATA_PROCESSED / "ticket_decomp.csv", index_col=0)
    el = pd.read_csv(DATA_PROCESSED / "elasticity.csv").iloc[0]
    clus = pd.read_csv(DATA_PROCESSED / "market_clusters.csv", index_col=0)
    dow = pd.read_csv(DATA_PROCESSED / "sales_dow.csv", index_col=0)
    ash = pd.read_csv(DATA_PROCESSED / "pop_age_shift.csv", index_col=0)
    dq = pd.read_csv(DATA_PROCESSED / "vacancy_data_quality.csv")
    av = pd.read_csv(DATA_PROCESSED / "area_vs_market.csv", index_col=0)
    elf = pd.read_csv(DATA_PROCESSED / "elasticity_final.csv")

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Cm(2.0)
    sec.left_margin = sec.right_margin = Cm(2.0)

    # ================= 표지 =================
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    set_font(p.add_run("MARKET INTELLIGENCE REPORT"), 9, True, ACCENT)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    set_font(p.add_run("청량리 전통시장 클러스터\n상권분석 및 Pain Point 진단"), 23, True, ACCENT)
    para(doc, "서울 동대문구 제기동·청량리동·용신동 일대 9개 전통시장", 11, MUTED, 2)
    para(doc, "2026년 8월", 11, MUTED, 16)

    table(doc, ["구분", "내용"], [
        ["대상", "9개 전통시장 · 점포 2,588개 · 상인 6,766명"],
        ["기간", "매출 2021~2025년(20개 분기), 생활인구 2017~2026년(110개월), 공실 2006~2023년"],
        ["자료", "서울시 상권분석서비스, 서울시 생활인구, 소상공인시장진흥공단, 국세청 상가업소정보"],
        ["처리 규모", "생활인구 원시 레코드 5,430만 행, 원본 138GB, 전국 1,812개 시장 교차검증"],
        ["기법", "변이-할당 분석, 고정효과 패널회귀, 계층적 군집분석, 매출 요인분해, 데이터 품질검증"],
    ], widths=[2.8, 13.6], cap="[표 1] 분석 개요")
    doc.add_page_break()

    # ================= 진단 요약 =================
    heading(doc, "진단 요약")

    table(doc, ["구분", "관측", "해석", "대응 방향"], [
        ["상권 위치", f"9개 중 8개가 서울 전통시장 283개 중 상위 20% 이내\n서울 중위 성장률 {el.seoul_median:.1f}% 대비 44~195%",
         "클러스터 전체의 쇠퇴는 확인되지 않음", "일괄 지원보다 시장별 선별 대응"],
        ["성장의 질", "9개 중 8개에서 객단가 하락\n동서시장 건수 +330% / 객단가 −44%",
         "물량 증가가 단가 하락을 상쇄한 구조", "고부가 품목·서비스 발굴로 단가 방어"],
        ["배후 수요", "제기동 야간인구 −7.8%, 60대 이상 −15.1%(약 1,237명)\n2023년 입주분은 전농1동·용신동에 귀착",
         "시장 도보권 밖에 주거가 공급됨", "정비사업 시 시장 인접지 주거 비중 확보"],
        ["인구-매출 관계", "서울 279개 시장 5,444개 관측치에서 계수 0.00\n신뢰구간 [−0.15, 0.15]",
         "인구 감소가 매출 감소로 이어진다는 전제는 미확인",
         "인구 지표를 근거로 한 지원 규모 산정 재검토"],
        ["서울약령시장", "매출 −11.6%, 점포 −59개\n업종구성효과 +16.5% / 경쟁력효과 −52.9%",
         "업종이 아니라 입지 경쟁력 문제", "타 시장과 분리한 업종 전환·구조 재편"],
        ["집객 사업", "경동시장 20~30대 매출 비중 5.3% → 3.7%",
         "방문객 증가가 고객 구성으로 전환되지 않음", "성과 지표에 매출 구성·점포 생존율 포함"],
        ["진단 지표", f"공실 통계 전년 동일 기록 {dq.iloc[-1]['동일비율%']:.1f}%",
         "현행 공실률은 정책 근거로 사용 불가", "카드매출·생활인구 기반 지표 병행"],
    ], widths=[2.2, 5.4, 3.9, 4.9], cap="[표 2] 관측과 대응")
    doc.add_page_break()

    # ================= 1. 서울 대비 위치 =================
    heading(doc, "1. 서울 전체 전통시장 대비 위치")
    para(doc, "서울시 상권분석서비스 등록 전통시장 상권 283개의 2021~2025년 매출 성장률 분포에서 "
              "청량리 9개 시장의 위치를 산출하였다.")

    rows = [[r["name"], f"{r.growth:,.1f}%", f"상위 {100 - r['백분위']:.0f}%"]
            for _, r in bm[bm["청량리"]].sort_values("growth", ascending=False).iterrows()]
    rows.append(["서울 전통시장 중위값", f"{el.seoul_median:,.1f}%", "—"])
    table(doc, ["시장", "매출 성장률", "서울 내 순위"], rows,
          widths=[7.0, 4.7, 4.7], align_right=(1,), align_center=(2,), cap="[표 3] 서울 전통시장 283개 중 청량리 9개의 위치")
    add_chart(doc, C.hbar(
        [(r["name"], float(r["백분위"])) for _, r in
         bm[bm["청량리"]].sort_values("백분위").iterrows()],
        unit="%", label="서울 전통시장 내 백분위"), "benchmark", cap="[그림 1] 서울 전통시장 283개 대비 백분위")

    para(doc, "청량리수산시장은 상위 1%, 동서시장은 상위 3%에 해당한다. "
              "서울약령시장만 하위 25% 구간에 위치하여 클러스터 내에서 예외적이다.")

    # ================= 2. 매출 구조 =================
    heading(doc, "2. 매출 증가의 구성")
    para(doc, "매출을 거래 건수와 객단가로 분해하였다.")

    rows = [[m, f"{r['매출증가%']:+,.1f}%", f"{r['건수증가%']:+,.1f}%", f"{r['객단가증가%']:+,.1f}%",
             f"{r.객단가21:,.0f}원", f"{r.객단가25:,.0f}원"]
            for m, r in dt.sort_values("객단가증가%").iterrows()]
    table(doc, ["시장", "매출", "거래 건수", "객단가", "객단가 2021", "객단가 2025"], rows,
          widths=[4.2, 2.4, 2.4, 2.4, 2.5, 2.5], align_right=(1, 2, 3, 4, 5), cap="[표 4] 매출 변화의 요인 분해 (2021→2025)")
    add_chart(doc, C.scatter(
        [(m, float(r["건수증가%"]), float(r["객단가증가%"]),
          C.WARN if r["객단가증가%"] < -30 else C.ACCENT) for m, r in dt.iterrows()],
        xlab="거래 건수 증가율 (%)", ylab="객단가 증가율 (%)",
        quadrant=(0, 0), height=250, label="건수와 객단가"), "ticket", cap="[그림 2] 거래 건수와 객단가의 관계")

    para(doc, "9개 중 8개 시장에서 객단가가 하락하였다. 매출 총액 증가는 거래 건수 증가에 기인하며, "
              "거래 한 건당 금액은 감소하였다.")

    table(doc, ["원인", "내용", "판단 근거"], [
        ["결제수단 전환", "현금으로 처리되던 소액 거래가 카드 기록에 포착",
         "간편결제 확산 시기와 일치. 전 시장 공통 현상인 점이 부합"],
        ["거래 구조 변화", "고액 도매 거래가 줄고 소액 소매로 대체",
         "서울약령시장은 건수 +28%에도 매출 −11.6%로 단가 하락폭이 더 큼"],
    ], widths=[3.0, 6.4, 7.0], cap="[표 5] 객단가 하락의 가능한 원인")
    note(doc, "두 원인의 기여도를 분리하려면 결제 건당 금액 분포 자료가 필요하며, "
              "현행 공개 자료로는 확인되지 않는다.")

    doc.add_page_break()

    # ================= 3. 변이-할당 =================
    heading(doc, "3. 부진 원인 분해: 업종 구성 대 입지 경쟁력")
    para(doc, "매출 변화를 서울 전통시장 전체 성장분, 보유 업종 구성에서 오는 몫, "
              "같은 업종 내 상대 성과로 분해하였다.")

    rows = [[m, f"{r.SE_억:,.0f}", f"{r.IM_억:,.0f}", f"{r.CE_억:,.0f}",
             f"{r['실제증감_억']:,.0f}", f"{r['CE_기여%']:+,.1f}%"]
            for m, r in ss.iterrows()]
    table(doc, ["시장", "전체 성장분", "업종 구성 효과", "경쟁력 효과", "실제 증감", "경쟁력 기여"],
          rows, widths=[3.8, 2.5, 2.8, 2.5, 2.4, 2.4], align_right=(1, 2, 3, 4, 5), cap="[표 6] 변이-할당 분석 (단위: 억원)")
    add_chart(doc, C.diverging(
        [(m, float(r["CE_기여%"])) for m, r in ss.iterrows()],
        unit="%", label="경쟁력 효과"), "shiftshare", cap="[그림 3] 시장별 경쟁력 효과 (기준 규모 대비)")

    para(doc, "서울약령시장의 업종 구성 효과는 +16.5%로 보유 업종이 서울 평균보다 성장하는 분야에 "
              "속한다. 경쟁력 효과 −52.9%가 이를 상쇄하여 매출은 339억 원 감소하였다.")

    table(doc, ["업종", "매출 비중", "약령시", "서울 전통시장", "격차"], [
        ["청과상", "68.0%", "−26.7%", "+52.1%", "**−78.8%p**"],
        ["의약품", "22.2%", "−0.3%", "+19.1%", "−19.4%p"],
        ["한의원", "4.9%", "−27.4%", "−5.4%", "−22.0%p"],
        ["슈퍼마켓", "2.9%", "−2.1%", "+5.2%", "−7.3%p"],
    ], widths=[3.4, 3.0, 3.2, 3.6, 3.2], align_right=(1, 2, 3, 4), cap="[표 7] 서울약령시장 주요 업종의 성장률 비교")
    para(doc, "약령시 매출의 68%는 한약재가 아닌 청과상에서 발생한다. 서울 전통시장 청과상이 52% "
              "성장한 기간에 약령시 구역은 27% 감소하여 격차가 79%p에 달한다. "
              "주요 4개 업종 모두 서울 평균을 밑돈다.")
    para(doc, "따라서 한방 수요 축소는 부진의 일부 요인일 뿐이며, 집객 중심 활성화 사업으로는 "
              "회복이 어렵다. 업종 전환과 상권 구조 재편을 별도 과제로 다루어야 한다.")

    doc.add_page_break()

    # ================= 4. 배후인구 =================
    heading(doc, "4. 배후인구 변화")
    para(doc, "시장 소재 행정동의 야간 생활인구(03~05시)를 상주인구 대리지표로 사용하였다.")

    rows = [[d, f"{ash.loc[d, '0-19_변화율']:+,.1f}%", f"{ash.loc[d, '20-39_변화율']:+,.1f}%",
             f"{ash.loc[d, '40-59_변화율']:+,.1f}%", f"{ash.loc[d, '60+_변화율']:+,.1f}%",
             "롯데캐슬 SKY-L65" if d == "전농1동" else
             ("한양수자인 그라시엘" if d == "용신동" else
              ("경동시장·서울약령시" if d == "제기동" else "청량리종합시장 등 5개"))]
            for d in ["전농1동", "용신동", "청량리동", "제기동"]]
    table(doc, ["행정동", "0~19세", "20~39세", "40~59세", "60세 이상", "주요 시설"], rows,
          widths=[2.4, 2.2, 2.2, 2.2, 2.4, 5.0], align_right=(1, 2, 3, 4), cap="[표 8] 행정동별 야간 생활인구 변화 (2022→2025)")
    para(doc, "2023년 청량리역 일대 약 2,800세대 입주 효과는 전농1동과 용신동에서 확인된다. "
              "경동시장과 서울약령시가 위치한 제기동은 같은 기간 인구가 감소하였다.")
    para(doc, "제기동 감소분은 60대 이상에 집중되어 있다. 청량리 시장 매출에서 60대 이상 비중이 "
              "46~72%인 점을 고려하면 주 고객층이 도보권에서 이탈하는 중이다. "
              "신축이 들어선 용신동·전농1동에서도 60대 이상 증가폭은 다른 연령대보다 작다.")

    heading(doc, "배후인구 감소가 매출로 이어지는가", level=2)
    para(doc, "인구 감소가 매출 감소로 연결되는지 검정하였다. 청량리 9개만으로는 표본이 작아 "
              "서울시 전통시장 상권 전체로 확대하고, 서울 424개 행정동의 야간 생활인구를 "
              "상권 경계의 행정동코드로 연결하였다. 상호 독립적인 두 설정으로 추정하였다.")

    rows = [[r.설정, r.표본, f"{int(r.관측치):,}", f"{r.beta:.2f}",
             f"[{r.ci_lo:.2f}, {r.ci_hi:.2f}]", f"{r.r2:.3f}"]
            for _, r in elf.iterrows()]
    table(doc, ["설정", "표본", "관측치", "추정 계수", "95% 신뢰구간", "R2"], rows,
          widths=[3.4, 4.6, 1.9, 2.0, 2.6, 1.9], align_right=(2, 3, 5), align_center=(4,),
          cap="[표 9] 배후인구-매출 관계 추정")

    para(doc, "두 설정 모두 계수가 0과 구분되지 않는다. 분기 고정효과 설정은 관측치 5,444개에서 "
              "신뢰구간이 [-0.15, 0.15]로 좁아, 배후인구 변화가 매출에 미치는 영향이 크다는 "
              "가능성까지 배제된다.")

    table(doc, ["구분", "내용"], [
        ["확인된 사실", "제기동 야간인구 -7.8%, 60대 이상 -15.1%. 서울약령시장 매출 -11.6%"],
        ["확인되지 않은 것", "배후인구 감소가 매출 감소의 원인이라는 연결"],
        ["가능한 설명", "분기·연 단위 인구 변동폭이 4~6%로 작아 매출 변동에 묻힘. "
                     "전통시장 매출은 상주인구보다 유입 고객·물가·결제수단 변화에 더 크게 좌우"],
        ["해석상 유의", "장기 수준 차이(인구가 많은 지역의 시장이 큰가)는 고정효과에 흡수되므로 "
                     "이 분석의 대상이 아님"],
    ], widths=[3.2, 13.2], cap="[표 10] 검정 결과 정리")

    para(doc, "배후인구 감소는 관측된 사실이나, 그것이 매출 감소로 이어진다는 통상적 전제는 "
              "서울 전통시장 전체 자료에서 지지되지 않는다. 인구 감소를 근거로 매출 영향을 "
              "추정하거나 지원 규모를 산정하는 방식은 재검토가 필요하다.")

    doc.add_page_break()

    # ================= 5. 시장 유형 =================
    heading(doc, "5. 거래 패턴 기반 시장 유형")
    para(doc, "각 시장을 시간대 6개, 요일 7개, 연령대 6개로 구성된 19차원 매출 구성비 벡터로 "
              "표현하고 계층적 군집분석(Ward)을 수행하였다.")

    names = {1: "주말 소매형", 2: "평일 도매형", 3: "혼합 대형", 4: "근린 생필품형"}
    lit = {"A": "한약재 도매", "B": "새벽 청과 도매", "C": "근린 소매", "D": "관광·F&B"}
    rows = [[m, lit.get(r["문헌축"], "—"), names[int(r["군집_4"])],
             f"{dow.loc[m, '주말비중']:.1f}%",
             "일치" if (r["문헌축"] == "A" and int(r["군집_4"]) == 2) or
                      (r["문헌축"] == "C" and int(r["군집_4"]) == 4) else "불일치"]
            for m, r in clus.sort_values("군집_4").iterrows()]
    table(doc, ["시장", "통념상 분류", "데이터 기반 유형", "주말 비중", "대조"], rows,
          widths=[3.8, 3.2, 3.4, 2.6, 2.4], align_right=(3,), align_center=(4,), cap="[표 11] 통념상 분류와 거래 패턴 기반 분류")
    para(doc, "서울약령시장만 단독 군집으로 분리된다. 주말 매출 비중 19.3%, 일요일 5.4%로 "
              "도매 성격이 데이터로 확인되는 유일한 시장이다.")

    table(doc, ["시장", "통념", "실제 데이터"], [
        ["경동시장", "한약재 도매",
         "업종 구성 청과상 35.3%, 미용재료 18.4%, 수산물 12.5%. 청량리종합시장과 동일 군집"],
        ["청량리청과물시장", "새벽 청과 도매", "00~06시 매출 2.3%, 주말 비중 32.3%"],
        ["동서시장", "새벽 청과 도매", "00~06시 매출 1.5%, 주말 비중 33.6%"],
    ], widths=[3.4, 3.2, 9.8], cap="[표 12] 통념과 데이터가 어긋나는 사례")
    para(doc, "청량리를 한약재 도매 집적지로 보는 인식은 9개 중 1개 시장에만 해당한다. "
              "클러스터 단위 일괄 지원은 실제 상권 성격과 어긋나며, 유형별 분리 설계가 필요하다.")

    doc.add_page_break()

    # ================= 6. 통계 신뢰성 =================
    heading(doc, "6. 공실 통계의 신뢰성")
    para(doc, "전통시장 진단에 널리 쓰이는 공실률의 사용 가능 여부를 확인하기 위해 "
              "소상공인시장진흥공단 연도별 자료를 전국 1,812개 시장, 15개 연도로 검증하였다.")

    add_chart(doc, C.hbar(
        [(r["구간"], float(r["동일비율%"])) for _, r in dq.tail(6).iterrows()],
        unit="%", label="전년 동일 기록 비율",
        highlight={"2021→2022", "2022→2023"}), "dq", cap="[그림 4] 전년도와 영업점포·빈점포가 완전히 동일한 시장의 비율")

    table(doc, ["연도", "영업점포", "빈점포", "총점포", "공실률", "비고"], [
        ["2019", "70", "1", "71", "1.4%", ""],
        ["2020", "73", "30", "103", "29.1%", "분모 32 증가"],
        ["2021", "62", "41", "103", "39.8%", ""],
        ["2022", "65", "5", "**70**", "7.1%", "**분모 33 감소**"],
        ["2023", "61", "5", "66", "7.6%", ""],
    ], widths=[2.2, 2.6, 2.4, 2.4, 2.4, 4.4], align_right=(1, 2, 3, 4), cap="[표 13] 청량리전통시장 원자료")
    table(doc, ["항목", "관측", "영향"], [
        ["기록 갱신", f"최근 구간 전국 시장의 {dq.iloc[-1]['동일비율%']:.1f}%가 전년과 동일한 값",
         "경동시장 2020~2022년 3개 연도가 영업 362·빈점포 316으로 동일. "
         "인용되는 46.6%는 관측된 변화가 아님"],
        ["모집단 안정성", "매년 약 20%의 시장에서 총점포수 ±20% 이상 변동",
         "청량리전통시장 2022년 공실률 급락은 분모 33개 감소의 결과"],
        ["전국 추세", "2017년 8.42% → 2023년 11.75% 단조 상승",
         "특정 연도 급락을 뒷받침할 전국 추세 없음"],
    ], widths=[2.6, 6.2, 7.6], cap="[표 14] 검증 결과")
    para(doc, "공표 자료를 그대로 사용하면 경동시장 공실률 46.9%, 청량리전통시장 회복이라는 결론에 "
              "도달하나 두 가지 모두 사실이 아니다. 원자료 검증 없이는 정반대의 판단이 도출될 수 있다.")
    para(doc, "따라서 이 클러스터의 상권 진단에는 공실률 대신 카드매출과 생활인구를 기준 지표로 "
              "사용하고, 공실률은 소진공 조사표 원표를 확인한 이후에 보조 지표로 활용해야 한다.")

    doc.add_page_break()

    # ================= 7. 실행 과제 =================
    heading(doc, "7. 실행 과제")
    para(doc, "진단 결과를 대응 시점과 소관을 기준으로 정리하였다.")

    table(doc, ["순위", "과제", "근거", "소관", "시점"], [
        ["1", "정비사업 배후인구 설계 검토\n청량리6·7·8구역 주거 배치가 시장 도보권을 포함하는지 확인",
         "2023년 입주분이 전농1동·용신동에 귀착, 제기동 −7.8%", "서울시·동대문구", "즉시"],
        ["2", "서울약령시장 개별 진단\n업종 전환 가능성과 상권 구조 재편 방향 수립",
         "경쟁력 효과 −52.9%, 주요 4개 업종 전부 서울 평균 미달", "동대문구·상인회", "즉시"],
        ["3", "성과 지표 개편\n방문객 수 외 매출 구성·점포 생존율·객단가 포함",
         "경동시장 20~30대 비중 5.3%→3.7%", "사업 시행 주체", "차기 사업 설계 시"],
        ["4", "고령 고객 접근성 확보\n제기동 이탈 인구의 이동 경로 파악 및 원거리 접근 수단 검토",
         "제기동 60대 이상 −15.1%, 약 1,237명", "동대문구", "6개월 내"],
        ["5", "공실 통계 정합성 확인\n소진공 조사표 원표 및 점포 모집단 정의 확인",
         "전년 동일 기록 47.6%", "소진공·발주기관", "6개월 내"],
        ["6", "객단가 하락 원인 규명\n결제수단 전환분과 실제 거래구조 변화분 분리",
         "9개 중 8개 시장 객단가 하락", "분석 기관", "후속 과제"],
    ], widths=[1.2, 5.2, 4.6, 2.8, 2.6], align_center=(0, 4), cap="[표 15] 과제별 우선순위")
    doc.add_page_break()

    # ================= 8. 전제 검증 =================
    heading(doc, "8. 분석 전제의 공간 검증")
    para(doc, "매출 자료의 공간 단위인 상권 영역이 시장의 물리적 범위와 일치하는지 확인하였다. "
              "상권 경계 폴리곤(EPSG:5181)에 국세청 상가업소정보의 점포 좌표 17,303개를 "
              "공간 조인하여 상권 내부 점포를 직접 집계하였다.")

    rows = [[m, f"{int(r.상권내_점포):,}", f"{int(r.시장_등록점포):,}", f"{r.배율:.1f}배",
             f"{int(r['상권면적_㎡']):,}"]
            for m, r in av.sort_values("배율", ascending=False).iterrows()]
    table(doc, ["시장", "상권 내 점포(실측)", "시장 등록 점포", "배율", "상권 면적(㎡)"], rows,
          widths=[4.0, 3.4, 3.2, 2.4, 3.4], align_right=(1, 2, 3, 4),
          cap="[표 16] 상권 영역 내 실측 점포와 시장 등록 점포")

    para(doc, "상권 영역이 시장보다 넓어 매출이 과대 계상될 것이라는 우려는 확인되지 않았다. "
              "9개 중 7개에서 상권 내 실측 점포가 시장 등록 점포보다 적거나 비슷하다. "
              "경동시장 0.4배, 동서시장 0.2배로 오히려 상권 쪽이 좁다.")

    table(doc, ["자료", "서울약령시장 점포 수", "성격"], [
        ["소진공 시장 등록", "951", "시장 관리 주체가 신고한 점포"],
        ["상권분석서비스", "2,347", "상권 영역 내 사업자 집계"],
        ["상가업소정보 실측", "542", "폴리곤 내 좌표가 확인된 점포"],
    ], widths=[4.4, 4.2, 7.8], align_right=(1,),
        cap="[표 17] 출처별 점포 수 차이 (서울약령시장 사례)")

    para(doc, "동일 시장에 대해 자료별로 542개에서 2,347개까지 4배 넘게 벌어진다. "
              "매출 총액의 절대 수준을 자료 간에 비교할 때는 이 차이를 감안해야 하며, "
              "본 보고서가 사용한 성장률과 구성비는 동일 자료 내 비교이므로 영향을 받지 않는다.")

    heading(doc, "경동1960 매출 귀속", level=2)
    table(doc, ["확인 항목", "결과"], [
        ["경동시장 상권 폴리곤 내 스타벅스", "존재 확인 (업종분류: 카페)"],
        ["해석에 미치는 영향", "스타벅스 매출이 포함된 상태에서 20~30대 비중이 "
                            "5.3%에서 3.7%로 감소. 집객 성과가 상권 고객 구성으로 "
                            "전환되지 않았다는 결론이 유지됨"],
    ], widths=[5.0, 11.4], cap="[표 18] 스타벅스 경동1960의 상권 귀속 확인")

    heading(doc, "잔여 한계", level=2)
    table(doc, ["항목", "내용", "해소 방안"], [
        ["카드매출 범위", "현금·계좌 거래가 많은 도매 상권은 매출이 과소 반영",
         "생활인구 시간대 분포로 보완. 완전 해소에는 사업자 매출신고 자료 필요"],
        ["배후인구-매출 관계", "서울 전역 표본에서 관계가 확인되지 않음(4장). "
                            "인구 감소의 매출 영향 경로는 미규명 상태",
         "야간인구 대신 주간 유입인구, 도보권 단위 집계로 재검정. "
         "표본은 이미 5,444개로 확대하였으므로 지표 선택 문제로 판단"],
        ["공실 통계", "소진공 자료의 기록 갱신 미흡",
         "조사표 원표 및 점포 모집단 정의 확인 필요"],
    ], widths=[2.8, 6.4, 7.2], cap="[표 19] 잔여 한계와 해소 방안")
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"저장: {OUT}")
    print("  표 19개 · 차트 4개 · 캡션 하단 가운데 정렬")


if __name__ == "__main__":
    build()
